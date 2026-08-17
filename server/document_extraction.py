"""Extracts a normalized text representation from an uploaded document -
Excel, PDF, Word, or HTML. No AI, no data mutation - this is the plumbing
phase 4b will consume. See docs/superpowers/specs/
2026-08-15-document-upload-phase4a-design.md section 3.
"""
import csv
import io
from pathlib import Path

import fitz
import openpyxl
import pdfplumber
from bs4 import BeautifulSoup
from docx import Document

MAX_FILE_SIZE_BYTES = 20 * 1024 * 1024  # 20MB


class UnsupportedFormatError(Exception):
    """Raised for a file extension with no parser - message safe to show
    the user directly."""


class ExtractionError(Exception):
    """Raised when a recognized format fails to parse, or parses to no
    content - message safe to show the user directly, never a raw
    traceback."""


def _extract_xlsx(content_bytes):
    try:
        workbook = openpyxl.load_workbook(io.BytesIO(content_bytes), data_only=True)
    except Exception as exc:
        raise ExtractionError(f"Could not read Excel file: {exc}") from exc

    blocks = []
    for sheet in workbook.worksheets:
        rows = []
        for row in sheet.iter_rows(values_only=True):
            if all(cell is None for cell in row):
                continue
            rows.append(" | ".join("" if cell is None else str(cell) for cell in row))
        if rows:
            blocks.append(f"Sheet: {sheet.title}\n" + "\n".join(rows))
    return "\n\n".join(blocks)


def _extract_docx(content_bytes):
    try:
        document = Document(io.BytesIO(content_bytes))
    except Exception as exc:
        raise ExtractionError(f"Could not read Word document: {exc}") from exc

    parts = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())
    for table_index, table in enumerate(document.tables, start=1):
        rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
        if rows:
            parts.append(f"Table {table_index}\n" + "\n".join(rows))
    return "\n\n".join(parts)


def _extract_html(content_bytes):
    try:
        soup = BeautifulSoup(content_bytes, "html.parser")
    except Exception as exc:
        raise ExtractionError(f"Could not read HTML file: {exc}") from exc

    for tag in soup(["script", "style"]):
        tag.decompose()

    # Pull tables out of the tree before extracting prose text, so a
    # table's cell contents aren't duplicated between the prose block and
    # the table's own structured block below.
    tables = soup.find_all("table")
    for table in tables:
        table.extract()

    text = soup.get_text(separator="\n", strip=True)
    parts = [text] if text else []

    for table_index, table in enumerate(tables, start=1):
        rows = []
        for tr in table.find_all("tr"):
            cells = tr.find_all(["td", "th"])
            row_text = " | ".join(cell.get_text(strip=True) for cell in cells)
            if row_text:
                rows.append(row_text)
        if rows:
            parts.append(f"Table {table_index}\n" + "\n".join(rows))
    return "\n\n".join(parts)


def _extract_txt(content_bytes):
    try:
        text = content_bytes.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise ExtractionError(f"Could not read text file: {exc}") from exc
    return text.strip()


def _extract_csv(content_bytes):
    try:
        text = content_bytes.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise ExtractionError(f"Could not read CSV file: {exc}") from exc
    reader = csv.reader(io.StringIO(text))
    rows = [" | ".join(cell.strip() for cell in row) for row in reader]
    rows = [row for row in rows if row.strip(" |")]
    return "\n".join(rows)


def _extract_pdf(content_bytes):
    parts = []
    try:
        with fitz.open(stream=content_bytes, filetype="pdf") as doc:
            for page_index in range(len(doc)):
                page_text = doc[page_index].get_text().strip()
                if page_text:
                    parts.append(f"Page {page_index + 1} text\n{page_text}")
    except Exception as exc:
        raise ExtractionError(f"Could not read PDF file: {exc}") from exc

    try:
        with pdfplumber.open(io.BytesIO(content_bytes)) as pdf:
            for page_index, page in enumerate(pdf.pages, start=1):
                for table_index, table in enumerate(page.extract_tables() or [], start=1):
                    rows = [" | ".join((cell or "").strip() for cell in row) for row in table]
                    rows = [row for row in rows if row.strip(" |")]
                    if rows:
                        parts.append(f"Page {page_index} table {table_index}\n" + "\n".join(rows))
    except Exception as exc:
        raise ExtractionError(f"Could not extract tables from PDF file: {exc}") from exc

    return "\n\n".join(parts)


# Legacy .xls binary files can't actually be read by openpyxl (it only
# understands the modern zip-based .xlsx format), but routing them through
# the same parser still produces a clean ExtractionError rather than
# silently misreporting the extension as unsupported - a mislabeled export
# occasionally turns out to be readable anyway.
_EXTENSION_PARSERS = {
    ".xlsx": _extract_xlsx,
    ".xls": _extract_xlsx,
    ".docx": _extract_docx,
    ".html": _extract_html,
    ".htm": _extract_html,
    ".pdf": _extract_pdf,
    ".txt": _extract_txt,
    ".csv": _extract_csv,
}


class ExtractionResult:
    def __init__(self, filename, format, text):
        self.filename = filename
        self.format = format
        self.text = text

    def to_dict(self):
        return {"filename": self.filename, "format": self.format, "text": self.text}


def extract(filename, content_bytes):
    if len(content_bytes) > MAX_FILE_SIZE_BYTES:
        raise ExtractionError(
            f"File is too large ({len(content_bytes):,} bytes) - the limit is {MAX_FILE_SIZE_BYTES:,} bytes."
        )
    extension = Path(filename).suffix.lower()
    parser = _EXTENSION_PARSERS.get(extension)
    if parser is None:
        raise UnsupportedFormatError(f"Unsupported file type: {extension or '(no extension)'}")
    text = parser(content_bytes)
    if not text.strip():
        raise ExtractionError("No extractable content found in this file.")
    return ExtractionResult(filename=filename, format=extension.lstrip("."), text=text)
