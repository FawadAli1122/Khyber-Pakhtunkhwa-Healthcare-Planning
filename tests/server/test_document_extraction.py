"""Unit tests for server/document_extraction.py. Every fixture is built
in-memory - no checked-in binary files. See docs/superpowers/specs/
2026-08-15-document-upload-phase4a-design.md section 5.
"""
import io

import fitz
import openpyxl
import pytest
from docx import Document

from server import document_extraction


def _build_xlsx_bytes():
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Districts"
    sheet.append(["district", "population"])
    sheet.append(["Peshawar", 4750388])
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def _build_docx_bytes():
    document = Document()
    document.add_paragraph("Peshawar has a population of 4,750,388.")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "district"
    table.rows[0].cells[1].text = "population"
    table.rows[1].cells[0].text = "Peshawar"
    table.rows[1].cells[1].text = "4750388"
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_pdf_bytes():
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Peshawar has a population of 4,750,388.")
    data = doc.tobytes()
    doc.close()
    return data


def _build_html_bytes():
    html = (
        "<html><body><p>Peshawar has a population of 4,750,388.</p>"
        "<table><tr><th>district</th><th>population</th></tr>"
        "<tr><td>Peshawar</td><td>4750388</td></tr></table>"
        "</body></html>"
    )
    return html.encode("utf-8")


def test_extract_xlsx_covers_every_sheet():
    result = document_extraction.extract("stats.xlsx", _build_xlsx_bytes())
    assert result.format == "xlsx"
    assert "Sheet: Districts" in result.text
    assert "Peshawar" in result.text
    assert "4750388" in result.text


def test_extract_docx_includes_paragraphs_and_tables():
    result = document_extraction.extract("notes.docx", _build_docx_bytes())
    assert result.format == "docx"
    assert "Peshawar has a population of 4,750,388." in result.text
    assert "Peshawar | 4750388" in result.text


def test_extract_pdf_includes_page_text():
    result = document_extraction.extract("report.pdf", _build_pdf_bytes())
    assert result.format == "pdf"
    assert "Peshawar has a population of 4,750,388." in result.text


def test_extract_html_includes_prose_and_table():
    result = document_extraction.extract("page.html", _build_html_bytes())
    assert result.format == "html"
    assert "Peshawar has a population of 4,750,388." in result.text
    assert "Peshawar | 4750388" in result.text


def test_extract_unsupported_extension_raises():
    with pytest.raises(document_extraction.UnsupportedFormatError):
        document_extraction.extract("data.json", b'{"district": "Peshawar"}')


def test_extract_corrupt_xlsx_raises_extraction_error():
    with pytest.raises(document_extraction.ExtractionError):
        document_extraction.extract("broken.xlsx", b"not a real xlsx file")


def test_extract_oversized_file_rejected_before_parsing():
    oversized = b"0" * (document_extraction.MAX_FILE_SIZE_BYTES + 1)
    with pytest.raises(document_extraction.ExtractionError):
        document_extraction.extract("huge.xlsx", oversized)


def test_extract_empty_document_raises_extraction_error():
    empty_html = b"<html><body></body></html>"
    with pytest.raises(document_extraction.ExtractionError):
        document_extraction.extract("empty.html", empty_html)


def test_extract_txt_returns_plain_text():
    result = document_extraction.extract("notes.txt", b"Peshawar has a population of 4,750,388.")
    assert result.format == "txt"
    assert "Peshawar has a population of 4,750,388." in result.text


def test_extract_csv_renders_pipe_delimited_rows():
    csv_bytes = b"district,population\nPeshawar,4750388\n"
    result = document_extraction.extract("stats.csv", csv_bytes)
    assert result.format == "csv"
    assert "district | population" in result.text
    assert "Peshawar | 4750388" in result.text


def test_extract_txt_undecodable_bytes_raises_extraction_error():
    with pytest.raises(document_extraction.ExtractionError):
        document_extraction.extract("bad.txt", b"\xff\xfe\x00\xff not valid utf-8 \xff")


def test_extract_csv_strips_utf8_bom():
    bom = b"\xef\xbb\xbf"
    csv_bytes = bom + b"district,population\nPeshawar,4750388\n"
    result = document_extraction.extract("stats.csv", csv_bytes)
    assert "﻿" not in result.text
    assert "district | population" in result.text
    assert "Peshawar | 4750388" in result.text


def test_extract_txt_strips_utf8_bom():
    bom = b"\xef\xbb\xbf"
    txt_bytes = bom + b"Peshawar has a population of 4,750,388."
    result = document_extraction.extract("notes.txt", txt_bytes)
    assert "﻿" not in result.text
    assert result.text == "Peshawar has a population of 4,750,388."
